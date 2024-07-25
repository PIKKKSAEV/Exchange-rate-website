from fastapi import FastAPI, HTTPException, Query, File, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse
import httpx
import logging
from datetime import datetime, timedelta
import databases, sqlalchemy
import aiofiles
import yfinance as yf
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from fastapi.staticfiles import StaticFiles
import openpyxl
import dropbox

app = FastAPI()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

DATABASE_URL = "Ваша ссылка на базу данных"
database = databases.Database(DATABASE_URL)
metadata = sqlalchemy.MetaData()

engine = sqlalchemy.create_engine(DATABASE_URL)
metadata.create_all(engine)

EXCHANGE_RATE_API_URL = "https://v6.exchangerate-api.com/v6/48e6c08df5b1329c0c0296c4/latest/USD"
TARGET_CURRENCIES = ["RUB", "EUR", "GBP", "CNY", "JPY"]
UPLOAD_DIR = "/Users/mihailpiksaev/Desktop/Летняя практика/uploads"

if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

DROPBOX_ACCESS_TOKEN = "Access token вашего DropBox"

@app.on_event("startup")
async def startup():
    await database.connect()
    scheduler = AsyncIOScheduler()
    scheduler.add_job(periodic_data_export, 'interval', hours=1)
    scheduler.start()

@app.on_event("shutdown")
async def shutdown():
    await database.disconnect()

@app.get("/", response_class=HTMLResponse)
async def read_root():
    async with aiofiles.open('static/index.html', mode='r') as f:
        html_content = await f.read()
    return HTMLResponse(content=html_content)

@app.get("/exchange-rates")
async def get_exchange_rates():
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(EXCHANGE_RATE_API_URL)
            response.raise_for_status()
            data = response.json()
            filtered_rates = {currency: rate for currency, rate in data["conversion_rates"].items() if currency in TARGET_CURRENCIES}
            return {"conversion_rates": filtered_rates}
    except httpx.HTTPStatusError as e:
        logger.error(f"HTTP error occurred: {e}")
        raise HTTPException(status_code=e.response.status_code, detail=f"Не удалось получить курсы валют: {e}")
    except httpx.RequestError as e:
        logger.error(f"Request error occurred: {e}")
        raise HTTPException(status_code=500, detail=f"Не удалось получить курсы валют из-за сетевой ошибки: {e}")

def get_yfinance_currency_symbol(currency: str) -> str:
    symbols = {
        "RUB": "RUB=X",
        "EUR": "EURUSD=X",
        "GBP": "GBPUSD=X",
        "CNY": "CNY=X",
        "JPY": "JPY=X"
    }
    return symbols.get(currency, None)

@app.get("/currency-history")
async def get_currency_history(currency: str = Query(...), days: int = Query(7)):
    symbol = get_yfinance_currency_symbol(currency)
    if not symbol:
        raise HTTPException(status_code=400, detail=f"Неверная валюта: {currency}")

    end_date = datetime.now()
    start_date = end_date - timedelta(days=days)
    try:
        ticker = yf.Ticker(symbol)
        hist = ticker.history(start=start_date.strftime('%Y-%m-%d'), end=end_date.strftime('%Y-%m-%d'))
        dates = hist.index.strftime('%Y-%m-%d').tolist()
        rates = hist['Close'].tolist()
        return {"dates": dates, "rates": rates}
    except Exception as e:
        logger.error(f"Ошибка получения исторических данных: {e}")
        raise HTTPException(status_code=500, detail=f"Не удалось получить исторические данные для {currency}: {e}")

@app.get("/generate-docx")
async def generate_docx():
    data = await get_exchange_rates()
    document = Document()
    document.add_heading('Exchange Rates', 0)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Currency'
    hdr_cells[1].text = 'Rate'

    for currency, rate in data['conversion_rates'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = currency
        row_cells[1].text = str(rate)

    docx_io = BytesIO()
    document.save(docx_io)
    docx_io.seek(0)

    headers = {
        'Content-Disposition': 'attachment; filename="exchange_rates.docx"'
    }
    return StreamingResponse(docx_io, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)

@app.get("/generate-pdf")
async def generate_pdf():
    data = await get_exchange_rates()
    pdf_io = BytesIO()
    c = canvas.Canvas(pdf_io, pagesize=letter)
    width, height = letter

    c.drawString(100, height - 40, "Exchange Rates")
    c.drawString(100, height - 60, "Currency")
    c.drawString(300, height - 60, "Rate")

    y = height - 80
    for currency, rate in data['conversion_rates'].items():
        c.drawString(100, y, currency)
        c.drawString(300, y, str(rate))
        y -= 20

    c.save()
    pdf_io.seek(0)

    headers = {
        'Content-Disposition': 'attachment; filename="exchange_rates.pdf"'
    }
    return StreamingResponse(pdf_io, media_type='application/pdf', headers=headers)

async def upload_to_dropbox(file_path: str, dropbox_path: str):
    try:
        dbx = dropbox.Dropbox(DROPBOX_ACCESS_TOKEN)
        with open(file_path, "rb") as f:
            dbx.files_upload(f.read(), dropbox_path)
        logger.info(f"File {file_path} uploaded successfully to Dropbox.")
    except Exception as e:
        logger.error(f"Failed to upload {file_path} to Dropbox: {e}")

async def periodic_data_export():
    try:
        logger.info("Starting periodic data export")
        data = await get_exchange_rates()
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

        document = Document()
        document.add_heading('Exchange Rates', 0)
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Currency'
        hdr_cells[1].text = 'Rate'
        for currency, rate in data['conversion_rates'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = currency
            row_cells[1].text = str(rate)
        docx_filename = os.path.join(UPLOAD_DIR, f"exchange_rates_{timestamp}.docx")
        document.save(docx_filename)
        
        await upload_to_dropbox(docx_filename, f"/exchange_rates_{timestamp}.docx")

        pdf_filename = os.path.join(UPLOAD_DIR, f"exchange_rates_{timestamp}.pdf")
        c = canvas.Canvas(pdf_filename, pagesize=letter)
        width, height = letter
        c.drawString(100, height - 40, "Exchange Rates")
        c.drawString(100, height - 60, "Currency")
        c.drawString(300, height - 60, "Rate")
        y = height - 80
        for currency, rate in data['conversion_rates'].items():
            c.drawString(100, y, currency)
            c.drawString(300, y, str(rate))
            y -= 20
        c.save()

        await upload_to_dropbox(pdf_filename, f"/exchange_rates_{timestamp}.pdf")

        logger.info(f"Data exported successfully: {docx_filename}, {pdf_filename}")

    except Exception as e:
        logger.error(f"Failed to export data: {e}")

@app.get("/upload", response_class=HTMLResponse)
async def upload_form():
    async with aiofiles.open('static/upload.html', mode='r') as f:
        html_content = await f.read()
    return HTMLResponse(content=html_content)

@app.post("/upload-xlsx/")
async def upload_xlsx(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        wb = openpyxl.load_workbook(BytesIO(contents))
        sheet = wb.active
        data = {}
        for row in sheet.iter_rows(min_row=2, values_only=True):
            currency, rate = row
            data[currency] = rate
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        xlsx_filename = os.path.join(UPLOAD_DIR, f"{timestamp}_{file.filename}")
        with open(xlsx_filename, 'wb') as f:
            f.write(contents)
        return {"filename": xlsx_filename, "data": data}
    except Exception as e:
        logger.error(f"Failed to upload XLSX file: {e}")
        raise HTTPException(status_code=500, detail=f"Не удалось обработать XLSX файл: {e}")

app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")
app.mount("/static", StaticFiles(directory="static"), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
